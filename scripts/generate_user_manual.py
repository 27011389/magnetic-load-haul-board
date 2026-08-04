from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Magnetic_Shiftboard_User_Manual.docx"
LOGO = ROOT / "public" / "ConsminLogo.png"
MANUAL_ASSETS = ROOT / "docs" / "manual-assets"

INK = "17241F"
GREEN = "526F43"
LIME = "D8FA4C"
PALE = "EEF3E0"
LIGHT = "F2F4F1"
MID = "68736F"
RED = "B9473E"
AMBER = "E7B04A"
VIOLET = "8063A8"
WHITE = "FFFFFF"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_borders(table, color="C7CDC8", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = Inches(width / 1440)
            tc_w = cell._tc.get_or_add_tcPr().first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def set_font(run, size=11, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_font(run, 9, color=MID)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    paragraph._p.append(fld)


def add_figure(path, width, caption, explanation, alt_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.keep_with_next = True
    picture = p.add_run().add_picture(str(path), width=Inches(width))
    picture._inline.docPr.set("title", caption)
    picture._inline.docPr.set("descr", alt_text)
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(3)
    cp.paragraph_format.keep_with_next = True
    set_font(cp.add_run(caption), 9, True, GREEN)
    ep = doc.add_paragraph()
    ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ep.paragraph_format.space_after = Pt(8)
    set_font(ep.add_run(explanation), 9, False, MID, True)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in (
    ("Heading 1", 16, GREEN, 18, 10),
    ("Heading 2", 13, GREEN, 14, 7),
    ("Heading 3", 12, INK, 10, 5),
):
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

for list_name in ("List Bullet", "List Number"):
    style = styles[list_name]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.paragraph_format.left_indent = Inches(0.375)
    style.paragraph_format.first_line_indent = Inches(-0.188)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.25

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = hp.add_run("MAGNETIC LOAD & HAUL SHIFTBOARD  |  USER MANUAL")
set_font(run, 8.5, True, MID)
footer = section.footer
add_page_number(footer.paragraphs[0])

# Editorial cover pattern.
if LOGO.exists():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(38)
    picture = p.add_run().add_picture(str(LOGO), width=Inches(2.5))
    picture._inline.docPr.set("title", "ConsMin logo")
    picture._inline.docPr.set("descr", "ConsMin company logo")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
r = p.add_run("OPERATOR REFERENCE GUIDE")
set_font(r, 10, True, GREEN)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(10)
r = p.add_run("Magnetic Load & Haul Shiftboard")
set_font(r, 28, True, INK)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(32)
r = p.add_run("Daily setup, allocation, fleet control and shift handover")
set_font(r, 14, False, MID)

table = doc.add_table(rows=1, cols=2)
set_table_geometry(table, [2700, 6660])
set_table_borders(table, "D7DDD8")
for idx, label in enumerate(("Document", "Details")):
    shade(table.rows[0].cells[idx], GREEN)
    set_font(table.rows[0].cells[idx].paragraphs[0].add_run(label), 10, True, WHITE)
set_repeat_table_header(table.rows[0])
metadata = [
    ("Audience", "Mine Control, supervisors, team leaders and shiftboard operators"),
    ("System", "ConsMin Magnetic Load & Haul Shiftboard"),
    ("Document type", "Operational instruction manual"),
    ("Revision", "1.0 | August 2026"),
]
for label, value in metadata:
    row = table.add_row()
    shade(row.cells[0], PALE)
    set_font(row.cells[0].paragraphs[0].add_run(label), 10, True, GREEN)
    set_font(row.cells[1].paragraphs[0].add_run(value), 10, False, INK)
    for cell in row.cells:
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
set_table_geometry(table, [2700, 6660])

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Use this board as an operational coordination aid. Confirm safety-critical assignments, competencies and equipment status through approved site systems and procedures.")
set_font(r, 9.5, False, RED, True)

doc.add_page_break()

doc.add_heading("Contents", level=1)
contents = [
    "1. Purpose and operating principles", "2. Mine 2, Mine 3 and Control operating model",
    "3. Board layout", "4. Starting a shift", "5. Working with magnets",
    "6. Allocating crews", "7. Truck and auxiliary fleet setup",
    "8. Pit allocation and operator linking", "9. Park-up and status lanes",
    "10. Board controls and display modes", "11. Shift handover procedure",
    "12. Troubleshooting", "13. Quick-reference checklists",
]
for item in contents:
    doc.add_paragraph(item, style="List Number")

doc.add_heading("1. Purpose and operating principles", level=1)
doc.add_paragraph("The shiftboard provides one shared visual view of Day and Night load-and-haul activity. It replaces loose physical magnets with draggable digital magnets that save to the shared board database.")
for text in [
    "Treat the board as a coordination display, not the authoritative maintenance, dispatch or training record.",
    "Confirm equipment availability, operator competency and work instructions before assignment.",
    "Keep the board unlocked only while making planned changes.",
    "Check the LIVE BOARD/SAVING indicator before leaving the screen.",
    "Use Undo immediately if a bulk action produces an unintended result.",
]:
    doc.add_paragraph(text, style="List Bullet")

doc.add_heading("2. Mine 2, Mine 3 and Control operating model", level=1)
doc.add_paragraph("The software presents the same shared board to every screen and does not enforce user roles. The responsibilities below are the recommended operating model: Control is the single editing authority, while Mine 2 and Mine 3 verify and consume the board for their areas. ICT access controls should support this model where practical.")
roles = [
    ("Control", "Own the live board", "Update shift details; allocate crews and fleet; resolve conflicts; confirm saves; lock the board; publish TV View."),
    ("Mine 2", "Verify Mine 2 activity", "Check the relevant pit, equipment, operator and park-up information; report discrepancies to Control; avoid concurrent edits unless Control hands over authority."),
    ("Mine 3", "Verify Mine 3 activity", "Check the relevant pit, equipment, operator and park-up information; report discrepancies to Control; avoid concurrent edits unless Control hands over authority."),
]
table = doc.add_table(rows=1, cols=3)
set_table_geometry(table, [1600, 2400, 5360])
set_table_borders(table)
for idx, label in enumerate(("Point", "Primary purpose", "Expected actions")):
    shade(table.rows[0].cells[idx], GREEN)
    set_font(table.rows[0].cells[idx].paragraphs[0].add_run(label), 10, True, WHITE)
set_repeat_table_header(table.rows[0])
for point, purpose, actions in roles:
    cells = table.add_row().cells
    set_font(cells[0].paragraphs[0].add_run(point), 9.5, True)
    set_font(cells[1].paragraphs[0].add_run(purpose), 9.5)
    set_font(cells[2].paragraphs[0].add_run(actions), 9.5)
set_table_geometry(table, [1600, 2400, 5360])

doc.add_heading("2.1 Control workflow", level=2)
for text in [
    "Confirm LIVE BOARD, unlock the board and set the date, roster and shift note.",
    "Choose the four- or five-section layout, load both crews and arrange fleet assignments.",
    "Verify active equipment, linked operators, unassigned warnings and all park-up/status lanes.",
    "Wait for LIVE BOARD after the final change, lock the board and select TV VIEW for shared displays.",
    "At handover, identify the next authorised editor and verbally confirm that editing control has transferred.",
]:
    doc.add_paragraph(text, style="List Number")

doc.add_heading("2.2 Mine 2 workflow", level=2)
for text in [
    "Open the board in TV View and select BOTH, DAYS or NIGHTS as required.",
    "Verify the Mine 2 work areas, equipment, operators and end-of-shift locations against the approved plan.",
    "Report a mismatch to Control with the shift, pit, unit/operator and expected correction.",
    "Confirm the correction appears and the screen returns to LIVE BOARD. Do not treat the display as the competency or maintenance authority.",
]:
    doc.add_paragraph(text, style="List Number")

doc.add_heading("2.3 Mine 3 workflow", level=2)
for text in [
    "Open the board in TV View and select BOTH, DAYS or NIGHTS as required.",
    "Verify the Mine 3 work areas, equipment, operators and end-of-shift locations against the approved plan.",
    "Report a mismatch to Control with the shift, pit, unit/operator and expected correction.",
    "Confirm the correction appears and the screen returns to LIVE BOARD. Do not treat the display as the competency or maintenance authority.",
]:
    doc.add_paragraph(text, style="List Number")

doc.add_heading("2.4 If Mine 2 or Mine 3 must edit", level=2)
doc.add_paragraph("Control must first stop editing and explicitly hand over editing authority. The receiving point exits TV View, confirms LIVE BOARD, makes the agreed change, waits for the save to complete, locks the board and advises Control. This reduces last-write-wins conflicts; it is an operating control, not a software-enforced lock between devices.")

doc.add_heading("3. Board layout", level=1)
doc.add_heading("3.1 Day and Night panels", level=2)
doc.add_paragraph("The board is divided evenly into Day Shift and Night Shift. Each side contains the same three columns:")
for text in [
    "Pit / Work Area - pit names, work areas and status locations.",
    "Asset / Supervisor - excavators, loaders, dozers, graders, water carts and supervisory magnets.",
    "Trucks / Operators - haul trucks and linked operator magnets.",
]:
    doc.add_paragraph(text, style="List Bullet")

add_figure(
    MANUAL_ASSETS / "board-overview.png",
    6.25,
    "Figure 1. Example Day Shift board arrangement",
    "The example shows four defined pit rows, compact truck stacks, the auxiliary reset block near the upper-middle, and crew magnets staged down the right edge.",
    "Example shiftboard showing four pit rows, trucks, auxiliary equipment and personnel magnets.",
)

doc.add_heading("3.2 Work sections", level=2)
doc.add_paragraph("Four pit sections are shown by default and span the full working area. Select + 5TH SECTION to add a fifth section. Select - 5TH SECTION to return to four. Magnets are proportionally repositioned when the section count changes.")

doc.add_heading("3.3 End-of-shift floor park-up strip", level=2)
doc.add_paragraph("Each pit section includes a dashed END-OF-SHIFT FLOOR PARK-UP / LV PICKUP strip. Two 20-pixel magnets fit above the strip; the larger area underneath is available for truck and operator arrangements.")

doc.add_heading("3.4 Bottom park-up lanes", level=2)
lanes = [
    ("GO LINE", "Ready equipment at Topvar, Radio Hill or Chris D go lines."),
    ("SHUT PAD", "Equipment parked at Radio Hill, Corgan, Chris D or Big Mack shut pads."),
    ("WORKSHOP", "Workshop go-line and workshop equipment."),
    ("STANDBY", "Unallocated or standby equipment."),
    ("GRAVEYARD", "Long-term park-up equipment."),
]
table = doc.add_table(rows=1, cols=2)
set_table_geometry(table, [2700, 6660])

add_figure(
    MANUAL_ASSETS / "park-up-lanes.png",
    6.3,
    "Figure 2. Bottom park-up and status lanes",
    "Read each lane from the coloured label at left, then place equipment fully inside the appropriate named zone. The count bubble at the right updates from magnet position.",
    "Bottom of the board showing go line, shut pad, workshop, standby and status lanes with count bubbles.",
)
set_table_borders(table)
table.rows[0].cells[0].text = "Lane"
table.rows[0].cells[1].text = "Use"
for cell in table.rows[0].cells:
    shade(cell, GREEN)
    for run in cell.paragraphs[0].runs:
        set_font(run, 10, True, WHITE)
set_repeat_table_header(table.rows[0])
for label, detail in lanes:
    cells = table.add_row().cells
    set_font(cells[0].paragraphs[0].add_run(label), 10, True, INK)
    set_font(cells[1].paragraphs[0].add_run(detail), 10, False, INK)
set_table_geometry(table, [2700, 6660])

doc.add_heading("4. Starting a shift", level=1)
doc.add_heading("4.1 Confirm connection", level=2)
doc.add_paragraph("Wait for LIVE BOARD before making changes. SAVING indicates a change is being written. RETRYING indicates the board cannot currently confirm the save.")
doc.add_heading("4.2 Set shift details", level=2)
for text in [
    "Select SHIFT / NOTE.",
    "Enter the board date and roster description.",
    "Update the shift note with the current operational instruction.",
    "Select Save. Confirm the header and note show the new information.",
]:
    doc.add_paragraph(text, style="List Number")

doc.add_heading("4.3 Choose the working layout", level=2)
doc.add_paragraph("Use four sections for the standard board. Add the fifth section only when another work area is required. Make this choice before arranging the shift where practical.")

doc.add_heading("5. Working with magnets", level=1)
doc.add_heading("5.1 Move a magnet", level=2)
for text in [
    "Unlock the board if BOARD LOCKED is displayed.",
    "Press and hold the magnet, then drag it to the required position.",
    "Magnets may pass over other magnets while dragging.",
    "Release only in a clear position. A red outline indicates an invalid drop; releasing there returns the magnet to its original position.",
]:
    doc.add_paragraph(text, style="List Number")

doc.add_heading("5.2 Precision movement", level=2)
doc.add_paragraph("SNAP 10PX aligns drag movements to the board grid. With a magnet selected, arrow keys move it one pixel; hold Shift with an arrow key to move ten pixels. Keyboard movement still prevents overlap.")

doc.add_heading("5.3 Add a reusable magnet", level=2)
for text in [
    "Select + MAGNET RACK.",
    "Filter by category or search for a unit or operator name.",
    "Click a rack magnet to add it, or drag it directly onto the board.",
    "Close the rack when finished to recover screen space.",
]:
    doc.add_paragraph(text, style="List Number")

doc.add_heading("5.4 Create or edit a magnet", level=2)
doc.add_paragraph("Select + CUSTOM to create a magnet. Double-click an existing magnet to edit it. The editor supports magnet type, main text, second line, width, height and colour. Person magnets also include Crew and Passed Out In / Competencies fields. Crew is retained in settings but is not printed on the magnet face.")

doc.add_heading("5.5 Duplicate or delete", level=2)
doc.add_paragraph("Open the magnet editor and choose DUPLICATE or DELETE. Deleting equipment safely unlinks any attached operator. Use Undo if the wrong magnet is removed.")

doc.add_heading("6. Allocating crews", level=1)
for text in [
    "Select ALLOCATE CREW.",
    "Choose A, B or C Crew and choose Days or Nights.",
    "Confirm the prompt. The selected crew is stacked down the right edge of that shift.",
    "Allocate the other shift independently. Allocating Nights does not move or remove the Day crew, and vice versa.",
    "Drag each operator from the crew rail to the required equipment.",
]:
    doc.add_paragraph(text, style="List Number")

doc.add_heading("6.1 Personnel display", level=2)
doc.add_paragraph("Crew magnets show first names only. Full names remain searchable and are retained internally. Double-click a person magnet to review or update crew and competency information.")

doc.add_heading("6.2 Clear all personnel", level=2)
doc.add_paragraph("CLEAR PERSONNEL removes every person magnet after confirmation while leaving assets, pits, notes and park-up lanes unchanged. The action can be reversed with Undo.")

doc.add_heading("7. Truck and auxiliary fleet setup", level=1)
doc.add_heading("7.1 Clean up trucks", level=2)
doc.add_paragraph("CLEAN UP TRUCKS ensures the complete truck fleet appears once on Days and once on Nights. Existing pit rows are retained where possible, missing trucks are distributed, and trucks are aligned into compact vertical stacks below the park-up strips. Linked operators move with their trucks.")

doc.add_heading("7.2 Reset auxiliary layout", level=2)
doc.add_paragraph("RESET AUX LAYOUT places the auxiliary fleet into two columns in the upper-middle of each shift:")
table = doc.add_table(rows=1, cols=2)
set_table_geometry(table, [4680, 4680])
set_table_borders(table)
for idx, label in enumerate(("Left column", "Right column")):
    shade(table.rows[0].cells[idx], GREEN)
    set_font(table.rows[0].cells[idx].paragraphs[0].add_run(label), 10, True, WHITE)
set_repeat_table_header(table.rows[0])
row = table.add_row().cells
set_font(row[0].paragraphs[0].add_run("DZ014, DZ017, DZ018, DZ019, WD001"), 10)
set_font(row[1].paragraphs[0].add_run("WC019, WC201, WC018, GR012, GR013, GR014"), 10)
set_table_geometry(table, [4680, 4680])
doc.add_paragraph("Excavators/diggers and light vehicles are excluded and remain in their current positions.")

add_figure(
    MANUAL_ASSETS / "auxiliary-fleet.png",
    3.4,
    "Figure 3. Auxiliary equipment included in the reset",
    "The auxiliary reset uses the graders, rubber-tyred dozers, track dozers and water tankers shown. The light vehicle is deliberately excluded.",
    "Equipment list showing graders, one light vehicle, rubber-tyred dozers, track dozers and water tankers.",
)

doc.add_heading("8. Pit allocation and operator linking", level=1)
doc.add_heading("8.1 Allocate equipment", level=2)
doc.add_paragraph("Move equipment into the appropriate Day or Night pit section. Position excavators and auxiliary assets in the asset column and trucks in the trucks/operators area. Use the darker horizontal boundaries to keep each pit visually separate.")

doc.add_heading("8.2 Link an operator", level=2)
for text in [
    "Drag a person magnet close to the intended equipment magnet.",
    "Release in a clear position beside the equipment.",
    "The board snaps the person to the nearest available equipment and marks both magnets LINK.",
    "Moving the equipment then moves the linked operator as a group.",
    "Dragging the operator independently detaches the previous link so it can be reassigned.",
]:
    doc.add_paragraph(text, style="List Number")

doc.add_heading("8.3 Unassigned warning", level=2)
doc.add_paragraph("The warning chip reports working-area person magnets that are not linked to equipment. Select the chip repeatedly to move through unassigned operators. Supervisors or intentionally unassigned personnel may still appear in this count depending on where their magnets are placed.")

doc.add_heading("8.4 Allocation totals", level=2)
doc.add_paragraph("The Day and Night ALLOC totals currently count trucks positioned anywhere inside the working area. A truck in a cleanup stack is therefore counted as allocated even if it has no operator or loading unit. Use the totals as a location count, not as proof of a complete operational assignment.")

doc.add_heading("9. Park-up and status lanes", level=1)
for text in [
    "Move equipment into the appropriate bottom lane at end of shift or when status changes.",
    "Use GO LINE for ready equipment, SHUT PAD for nominated shut pads, WORKSHOP for workshop status, STANDBY for unallocated equipment, and GRAVEYARD for long-term park-up.",
    "The number at the right of each zone updates automatically from equipment position.",
    "Keep magnets fully inside a zone so the count is unambiguous.",
]:
    doc.add_paragraph(text, style="List Number")

doc.add_heading("10. Board controls and display modes", level=1)
controls = [
    ("UNDO", "Restores the previous committed board change."),
    ("FIND", "Searches visible board magnets by unit, first name, full name or second line."),
    ("SAVE START", "Stores the current layout as the layout used by RESET."),
    ("RESET", "Returns every magnet to the saved starting layout, or the supplied default if no start has been saved."),
    ("COPY SHIFT", "Copies the complete Day work area to Night, or Night to Day, replacing the destination work area."),
    ("LOCK BOARD", "Prevents moving, editing, adding, deleting and bulk-changing magnets."),
    ("TV VIEW", "Shows only the board and scales it to the display."),
    ("FULL SCREEN", "Requests browser full-screen mode while retaining edit controls."),
]
table = doc.add_table(rows=1, cols=2)
set_table_geometry(table, [2200, 7160])

add_figure(
    MANUAL_ASSETS / "quick-controls.png",
    6.3,
    "Figure 4. Quick-action control strip and allocation totals",
    "Use the quick controls for crew allocation, personnel clearing, fleet cleanup, auxiliary reset, section count, saved starting layout, shift copying and full-screen display. The Day/Night totals at right are positional truck counts.",
    "Top control strip showing Find, Unassigned, Allocate Crew, Clear Personnel, Clean Up Trucks, Reset Aux Layout and other board controls.",
)
set_table_borders(table)
for idx, label in enumerate(("Control", "Purpose")):
    shade(table.rows[0].cells[idx], GREEN)
    set_font(table.rows[0].cells[idx].paragraphs[0].add_run(label), 10, True, WHITE)
set_repeat_table_header(table.rows[0])
for label, detail in controls:
    cells = table.add_row().cells
    set_font(cells[0].paragraphs[0].add_run(label), 9.5, True)
    set_font(cells[1].paragraphs[0].add_run(detail), 9.5)
set_table_geometry(table, [2200, 7160])

doc.add_heading("11. Shift handover procedure", level=1)
for text in [
    "Confirm the board date, roster and shift note.",
    "Confirm the correct crews are loaded on Days and Nights.",
    "Verify each active loading unit, truck and auxiliary assignment against the approved operational plan.",
    "Review the unassigned-person warning and resolve or explain exceptions.",
    "Move unavailable equipment to the correct shut pad, workshop, standby or graveyard lane.",
    "Check park-up counts and visually inspect each lane.",
    "Confirm the status indicator reads LIVE BOARD and no save is pending.",
    "Lock the board, then use TV VIEW for the operational display if required.",
]:
    doc.add_paragraph(text, style="List Number")

doc.add_heading("12. Troubleshooting", level=1)
issues = [
    ("Magnet returns after release", "It overlapped another magnet. Drop it in a clear position; red means invalid."),
    ("Operator will not link", "Move closer to equipment, confirm another operator is not already linked, and ensure the drop position is clear."),
    ("Cannot edit or move", "The board is locked or TV View is active. Exit TV View and unlock the board."),
    ("Change appears not to save", "Wait for LIVE BOARD. If RETRYING persists, preserve the current screen and contact support before refreshing."),
    ("Wrong bulk action", "Select UNDO immediately. Bulk actions are recorded in history."),
    ("Truck total shows 28 allocated", "Cleanup placed all trucks inside the work area. The total is positional, not proof of assignment."),
    ("Horizontal scrollbar", "It appears only when the display is narrower than the fixed board. Use TV View or Full Screen for a fitted display."),
    ("Missing magnet", "Use FIND, then check the rack. Full personnel names are searchable even when only first names are shown."),
]
table = doc.add_table(rows=1, cols=2)
set_table_geometry(table, [3000, 6360])
set_table_borders(table)
for idx, label in enumerate(("Symptom", "Action")):
    shade(table.rows[0].cells[idx], GREEN)
    set_font(table.rows[0].cells[idx].paragraphs[0].add_run(label), 10, True, WHITE)
set_repeat_table_header(table.rows[0])
for issue, action in issues:
    cells = table.add_row().cells
    set_font(cells[0].paragraphs[0].add_run(issue), 9.5, True)
    set_font(cells[1].paragraphs[0].add_run(action), 9.5)
set_table_geometry(table, [3000, 6360])

doc.add_heading("13. Quick-reference checklists", level=1)
doc.add_heading("Before editing", level=2)
for text in ["LIVE BOARD displayed", "Correct shift/date confirmed", "Board unlocked", "Four/five-section layout selected"]:
    doc.add_paragraph("[  ] " + text)
doc.add_heading("Before handover", level=2)
for text in ["Crews correct", "Equipment and operators verified", "Unassigned warning reviewed", "Park-up lanes checked", "Latest change saved", "Board locked"]:
    doc.add_paragraph("[  ] " + text)
doc.add_heading("Mine 2 / Mine 3 verification", level=2)
for text in ["Correct shift view selected", "Relevant pits and fleet checked", "Discrepancies reported to Control", "Correction visible", "LIVE BOARD displayed"]:
    doc.add_paragraph("[  ] " + text)

doc.add_heading("Support information", level=1)
doc.add_paragraph("When reporting a problem, record the date and time, shift, browser/device, status indicator, action attempted and affected magnet names. Include a screenshot where possible. Do not repeatedly refresh while SAVING or RETRYING without first recording the board state.")

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.core_properties.title = "Magnetic Load & Haul Shiftboard User Manual"
doc.core_properties.subject = "Operational instructions for the digital load and haul whiteboard"
doc.core_properties.author = "ConsMin Operations"
doc.core_properties.keywords = "shiftboard, whiteboard, load and haul, operator manual"
doc.save(OUT)
print(OUT)
